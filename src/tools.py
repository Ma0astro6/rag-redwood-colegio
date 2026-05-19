import os
from langchain_core.tools import tool
from langchain_mongodb import MongoDBAtlasVectorSearch
from pymongo import MongoClient
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import PyPDF2
import docx
import requests
from bs4 import BeautifulSoup

# 1. Conexión a la nube (¡PEGA TU LINK AQUÍ Y CAMBIA EL <password>!)
MONGO_URI = "mongodb+srv://matias:redwood2026@cluster0.baqowbh.mongodb.net/?appName=Cluster0"

# Nos conectamos al servidor y a la colección exacta
cliente_mongo = MongoClient(MONGO_URI)
coleccion_mongo = cliente_mongo["redwood_db"]["documentos"]

# Preparamos el modelo de embeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Enlazamos LangChain con MongoDB
vectorstore = MongoDBAtlasVectorSearch(
    collection=coleccion_mongo,
    embedding=embeddings,
    index_name="vector_index",
    text_key="text",           
    embedding_key="embedding"  
)

# 2. Carga Inteligente (Sube los archivos solo 1 vez)
try:
    if coleccion_mongo.count_documents({}) == 0:
        print("☁️ La base de datos está vacía. Subiendo documentos a la nube...")
        carpeta_public = "public"
        texto_completo = ""
        archivos_leidos = 0

        if os.path.exists(carpeta_public):
            for nombre_archivo in os.listdir(carpeta_public):
                ruta_archivo = os.path.join(carpeta_public, nombre_archivo)
                
                if nombre_archivo.endswith(".pdf"):
                    with open(ruta_archivo, "rb") as archivo:
                        lector = PyPDF2.PdfReader(archivo)
                        for pagina in lector.pages:
                            texto_completo += pagina.extract_text() or ""
                    archivos_leidos += 1
                    print(f"📄 Procesando PDF: {nombre_archivo}")
                    
                elif nombre_archivo.endswith(".docx") and not nombre_archivo.startswith("~$"):
                    doc = docx.Document(ruta_archivo)
                    
                    # 1. Extraemos los párrafos normales
                    for parrafo in doc.paragraphs:
                        if parrafo.text.strip():
                            texto_completo += parrafo.text + "\n"
                    
                    # 2. Extraemos la información de las TABLAS
                    for tabla in doc.tables:
                        for fila in tabla.rows:
                            # Juntamos cada celda de la fila separada por " | " para que la IA entienda que es una tabla
                            fila_texto = [celda.text.replace('\n', ' ').strip() for celda in fila.cells if celda.text.strip()]
                            if fila_texto:
                                texto_completo += " | ".join(fila_texto) + "\n"
                        texto_completo += "\n" # Espacio al terminar la tabla
                        
                    archivos_leidos += 1
                    print(f"📝 Procesando Word (con tablas): {nombre_archivo}")

            if archivos_leidos > 0:
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=800) 
                chunks = text_splitter.split_text(texto_completo)
                
                # ¡La magia! Se suben todos los vectores a MongoDB
                vectorstore.add_texts(chunks)
                print(f"✅ ¡Éxito! {archivos_leidos} documentos subidos y vectorizados en MongoDB.")
            else:
                print("⚠️ No hay documentos válidos en la carpeta public.")
        else:
            print("❌ No existe la carpeta public.")
    else:
        print("✅ Conexión exitosa a MongoDB. Los documentos ya estaban en la nube.")
except Exception as e:
    print(f"❌ Error de conexión con MongoDB: {e}")


# 3. LAS HERRAMIENTAS
@tool
def buscar_en_documentos(pregunta: str) -> str:
    """Úsala para buscar información sobre reglamentos, útiles, lecturas o profesores."""
    print(f"🚨 EL AGENTE ESTÁ BUSCANDO EN MONGO: {pregunta}") 
    docs = vectorstore.similarity_search(pregunta, k=3)
    contexto_encontrado = "\n\n".join(doc.page_content for doc in docs)
    
    # 👇 RAYO X: Vemos si Mongo realmente devuelve texto
    print(f"📦 MONGO DEVOLVIÓ {len(docs)} DOCUMENTOS CON {len(contexto_encontrado)} CARACTERES.")
    
    if len(contexto_encontrado) == 0:
        return "No se encontró información en la base de datos."
    return contexto_encontrado

@tool
def consultar_web_colegio() -> str:
    """Úsala para buscar teléfonos, correos electrónicos, dirección o contacto."""
    print("🚨 EL AGENTE ESTÁ LEYENDO LA PÁGINA WEB OFICIAL") 
    
    url = "https://www.redwoodcollege.cl/" 
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)',
            'Accept-Language': 'es-ES,es;q=0.9'
        }
        respuesta = requests.get(url, headers=headers, timeout=10)
        respuesta.raise_for_status()
        
        soup = BeautifulSoup(respuesta.text, 'html.parser')
        texto_limpio = soup.get_text(separator=' ', strip=True)
        
        # 👇 RAYO X: Vemos si la web devuelve texto
        print(f"🌐 LA WEB DEVOLVIÓ {len(texto_limpio)} CARACTERES.") 
        
        if len(texto_limpio) < 50:
            return "Error: Pude entrar a la web, pero el texto está oculto."
            
        return f"Información de la web: {texto_limpio[:3000]}"
    except Exception as e:
        error_msg = f"No se pudo acceder a la página web. Error: {e}"
        print(f"❌ ERROR WEB: {error_msg}")
        return error_msg